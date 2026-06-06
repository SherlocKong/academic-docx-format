#!/usr/bin/env python3
"""Create a read-only structural and package inventory of a DOCX file."""

from __future__ import annotations

import argparse
import json
import re
import zipfile
from collections import Counter
from pathlib import Path

from docx import Document


FIELD_RE = re.compile(r"\b(SEQ|REF|PAGEREF|TOC|PAGE|NUMPAGES)\b")


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--output", type=Path)
    args = parser.parse_args()

    path = args.docx.resolve()
    if path.suffix.lower() != ".docx" or not path.is_file():
        parser.error("input must be an existing .docx file")

    document = Document(path)
    style_counts = Counter(p.style.name if p.style else "(none)" for p in document.paragraphs)
    with zipfile.ZipFile(path) as archive:
        names = archive.namelist()
        xml_text = "\n".join(
            archive.read(name).decode("utf-8", errors="ignore")
            for name in names
            if name.endswith(".xml")
        )
    report = {
        "path": str(path),
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "sections": len(document.sections),
        "inline_shapes": len(document.inline_shapes),
        "paragraph_style_counts": dict(style_counts.most_common()),
        "package_parts": len(names),
        "media_parts": sum(name.startswith("word/media/") for name in names),
        "header_parts": sum(name.startswith("word/header") and name.endswith(".xml") for name in names),
        "footer_parts": sum(name.startswith("word/footer") and name.endswith(".xml") for name in names),
        "has_comments": "word/comments.xml" in names,
        "tracked_insertions": len(re.findall(r"<w:ins(?:\s|>)", xml_text)),
        "tracked_deletions": len(re.findall(r"<w:del(?:\s|>)", xml_text)),
        "omml_equations": len(re.findall(r"<m:oMath(?:\s|>)", xml_text)),
        "omml_equation_paragraphs": len(re.findall(r"<m:oMathPara(?:\s|>)", xml_text)),
        "bookmarks": xml_text.count("<w:bookmarkStart"),
        "field_tokens": dict(Counter(FIELD_RE.findall(xml_text))),
        "package_part_names": names,
    }
    rendered = json.dumps(report, indent=2, ensure_ascii=False)
    if args.output:
        args.output.write_text(rendered + "\n", encoding="utf-8")
    else:
        print(rendered)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
